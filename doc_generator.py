"""
doc_generator.py - Build a formatted DOCX report using python-docx.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import STATUS_COLORS_DOCX, image_exists, logger


# ---------------------------------------------------------------------------
# Layout constants
# ---------------------------------------------------------------------------
IMAGE_MAX_WIDTH  = Inches(5.5)   # cap image width so it fits A4 margins
SECTION_SPACING  = Pt(8)         # space after each activity block


def generate_docx(activities: list[dict], output_path: str) -> str:
    """
    Create a DOCX report for *activities* and save it to *output_path*.

    Returns the absolute path of the saved file.
    Raises RuntimeError on failure.
    """
    logger.info("Generating DOCX report → %s", output_path)

    try:
        doc = Document()
        _set_page_margins(doc)
        _add_document_title(doc)
        _add_summary_table(doc, activities)
        _add_page_break(doc)

        for item in activities:
            _add_activity_section(doc, item)

        doc.save(output_path)
        logger.info("DOCX saved: %s", output_path)
        return os.path.abspath(output_path)

    except Exception as exc:
        raise RuntimeError(f"Failed to generate DOCX: {exc}") from exc


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _set_page_margins(doc: Document):
    """Set comfortable A4 margins on all sections."""
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)


def _add_document_title(doc: Document):
    """Add the main report title and a horizontal rule."""
    title_para = doc.add_heading("Weekly Sanity Check Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)   # dark navy

    # Subtitle / date line
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Automated System Health Overview")
    run.font.size  = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.italic = True

    doc.add_paragraph()  # blank line


def _add_summary_table(doc: Document, activities: list[dict]):
    """Insert a compact summary table at the top of the document."""
    doc.add_heading("Executive Summary", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"

    # Header row
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ("S.No", "Activity", "Status")):
        cell.text = label
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for item in activities:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["sno"])
        row_cells[1].text = str(item["activity"])

        status = item.get("status", "N/A")
        status_cell = row_cells[2]
        status_cell.text = ""
        run = status_cell.paragraphs[0].add_run(status)
        run.bold = True
        rgb = STATUS_COLORS_DOCX.get(status, (0, 0, 0))
        run.font.color.rgb = RGBColor(*rgb)

    doc.add_paragraph()  # spacing after table


def _add_activity_section(doc: Document, item: dict):
    """Add a full section for a single activity (heading, body, image)."""
    status  = item.get("status", "N/A")
    title   = str(item.get("doc_title", item.get("activity", "Activity")))
    desc    = str(item.get("doc_description", "No description provided."))
    img_raw = item.get("image")

    # --- Section heading ---
    heading = doc.add_heading(title, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    # --- Status badge paragraph ---
    status_para = doc.add_paragraph()
    status_para.add_run("Status: ").bold = True
    status_run = status_para.add_run(status)
    status_run.bold = True
    rgb = STATUS_COLORS_DOCX.get(status, (0, 0, 0))
    status_run.font.color.rgb = RGBColor(*rgb)

    # --- Description ---
    desc_para = doc.add_paragraph(desc)
    desc_para.paragraph_format.space_after = SECTION_SPACING

    # --- Image (if provided and found) ---
    if img_raw:
        abs_img = image_exists(img_raw)
        if abs_img:
            try:
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic_para.add_run()
                run.add_picture(abs_img, width=IMAGE_MAX_WIDTH)
            except Exception as exc:
                logger.warning("Could not insert image '%s': %s", abs_img, exc)
        # else: already logged by image_exists()

    # Horizontal divider
    _add_horizontal_line(doc)
    doc.add_paragraph()   # breathing room


def _add_horizontal_line(doc: Document):
    """Insert a thin horizontal rule paragraph."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_break(doc: Document):
    """Insert a page break paragraph."""
    doc.add_page_break()
